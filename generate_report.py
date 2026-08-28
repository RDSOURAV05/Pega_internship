import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_report():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(100)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("Movie Ticket Request Application\n(Kerala Theatrical Catalog)")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Pega Navy
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(180)
    subtitle_run = subtitle_p.add_run("Pega National Internship Program Capstone Deliverable\nSystem Implementation & Verification Report")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Intern Name: R. D. Sourav\nCase Type Name: Movie Ticket Request\nUniversity: APJ Abdul Kalam Technological University")
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(11)
    meta_run.font.bold = True
    meta_run.font.color.rgb = RGBColor(0x2F, 0x3E, 0x46)
    
    doc.add_page_break()
    
    # Section 1: Executive Summary & Scaffold
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. Executive Summary & Blueprint Scaffold")
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    doc.add_paragraph(
        "This project report documents the implementation of the 'Movie Ticket Request' case type on the Pega Academy platform. "
        "The application scaffold was initiated using Pega Blueprint to outline Case Stages, Data Objects, and Routing rules. "
        "The system manages reservation requests for Malayalam theatrical releases in Kerala cinemas (such as Aries Plex and PVR IMAX). "
        "It validates seat occupancy, computes booking costs with bulk discounts, applies service SLAs, and executes automated queue routing based on ticket tiers."
    )
    
    # Section 2: Case Lifecycle Design
    h2 = doc.add_heading(level=1)
    run = h2.add_run("2. Case Lifecycle Stage Configuration")
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    doc.add_paragraph(
        "Following the Pega Academy exercise guidelines, the 'Movie Ticket Request' case type was structured "
        "with four core stages and specific automation checkpoints:"
    )
    
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Stage Name", "Process / Steps", "Configured Rules & SLA"]
    for i, head in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = head
        set_cell_background(cell, "1F4E79")
        set_cell_margins(cell, 150, 150, 150, 150)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    stages_data = [
        ("1. Initial Stage", "Submit Movie Ticket Request form (customer & movie fields)", "Collect Information Step, Movie & Show Data Objects"),
        ("2. Availability", "Check Show Availability (Seat grid visualizer)", "Data Page verification, Row categories validation"),
        ("3. Approval", "Calculate Booking Cost (Calculate properties & check)", "Declare Expressions, 18% Tax, 10% Bulk Discount"),
        ("4. Booking Execution", "Case Resolution & Confirmation (Ticket release)", "Email Correspondence, Case Status (Resolved-Completed)")
    ]
    
    for row_idx, (stage, steps, rules) in enumerate(stages_data, start=1):
        table.cell(row_idx, 0).text = stage
        table.cell(row_idx, 1).text = steps
        table.cell(row_idx, 2).text = rules
        for col_idx in range(3):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell, 100, 100, 150, 150)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F2F2F2")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 3: Detailed Technical Implementations
    h3 = doc.add_heading(level=1)
    run = h3.add_run("3. System Specifications & Implementations")
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    # Data Structures
    doc.add_heading("3.1 Data Structures (Movie & Show Data Objects)", level=2)
    doc.add_paragraph(
        "Two core Pega Data Objects were established to support this booking flow:\n"
        "• Movie Data Object: Stores metadata for available films, including Title, Cast, Genre, and Release Date.\n"
        "• Show Data Object: Stores operational variables, including Theater Location, Date, Showtime, Available Seats, and Pricing Tiers."
    )
    
    # Calculations
    doc.add_heading("3.2 Total Cost Calculated Properties", level=2)
    doc.add_paragraph(
        "Pricing logic was implemented in an automated calculation engine. Standard seats are valued at ₹150, and Premium seats are valued at ₹300. "
        "The system automatically computes the total cost using the following rule parameters:\n"
        "• Base Ticket Cost = Selected Seats x Category Rate\n"
        "• Service Tax / GST = Base Ticket Cost x 18%\n"
        "• Bulk Booking Discount = 10% of Base Ticket Cost (triggered only if Ticket Quantity >= 4)\n"
        "• Net Total Cost = Base Ticket Cost + Service Tax - Bulk Discount"
    )
    
    # SLA Logic
    doc.add_heading("3.3 Service Level Agreement (SLA) Configuration", level=2)
    doc.add_paragraph(
        "To ensure reservations are processed in a timely manner, an SLA rule was attached to the case type:\n"
        "• Goal: 1 Day (The booking must be calculated and approved within 24 hours of creation).\n"
        "• Deadline: 2 Days (Urgent processing actions are triggered if the case remains unresolved after 48 hours)."
    )
    
    # Work Queue Routing
    doc.add_heading("3.4 Work Queue Routing (Premium vs. Standard)", level=2)
    doc.add_paragraph(
        "To balance workflow tasks, routing rules direct cases based on Show Type:\n"
        "• Premium Tickets: Automatically routed to Premium ShowQueue for priority audit.\n"
        "• Standard Tickets: Automatically routed to Standard ShowQueue."
    )
    
    # Correspondence
    doc.add_heading("3.5 Email Correspondence Rules", level=2)
    doc.add_paragraph(
        "An email correspondence rule is triggered during the Booking Execution stage. "
        "Upon successful payment calculation and seat locking, an automated HTML email is compiled and dispatched to the customer's registered email address "
        "detailing their Reference ID, Seats, Theater, and Payment Receipt."
    )
    
    # Section 4: Verification Checkpoints
    doc.add_page_break()
    h4 = doc.add_heading(level=1)
    run = h4.add_run("4. Screenshots & Verification Checkpoints")
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    doc.add_paragraph(
        "The following sections contain the execution screenshots captured during system run verification:"
    )
    
    def add_screenshot_box(title, description, image_filename):
        p = doc.add_paragraph()
        run = p.add_run(f"■ {title}")
        run.font.bold = True
        run.font.size = Pt(12)
        
        doc.add_paragraph(description)
        
        box_table = doc.add_table(rows=1, cols=1)
        box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = box_table.cell(0, 0)
        set_cell_background(cell, "FAFAFA")
        set_cell_margins(cell, 200, 200, 200, 200) # tight margin for image
        
        cell_p = cell.paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        image_path = os.path.join("C:\\Users\\PRO\\OneDrive\\Documents\\GitHub\\pega\\screenshots", image_filename)
        if os.path.exists(image_path):
            run = cell_p.add_run()
            # Set width to 5.5 inches to fit nicely in 8.5-inch page with 1-inch margins
            run.add_picture(image_path, width=Inches(5.5))
        else:
            run = cell_p.add_run(f"\n\n\n[ Screenshot File Not Found: {image_filename} ]\n\n\n")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
    add_screenshot_box("Pega Case Lifecycle Flow (Initial, Availability, Approval, Booking Execution)", "Scaffolded workflow showing Pega case stages, routing nodes, and SLA configurations.", "1_case_lifecycle.png")
    add_screenshot_box("US-001: Initial Stage Movie Ticket Request UI", "Submit Movie Ticket Request form showing Customer details, select Kerala Theater, and selected movie poster card.", "2_submit_request.png")
    add_screenshot_box("US-002: Availability Stage Seating Map UI", "Check Show Availability seat map showing standard rows, premium rows, occupied seats, and selected seat nodes.", "3_seating_map.png")
    add_screenshot_box("US-003: Approval Stage Cost Calculation UI", "Receipt billing invoice showing Base Fare, 18% Tax, 10% Bulk Discount deduction, and total net payable.", "4_cost_calculation.png")
    add_screenshot_box("Booking Execution Resolved Case Confirmation", "Resolved-Completed case page displaying transaction reference ID, SLA timelines, and case status transition log.", "5_case_resolution.png")
    
    # Save document
    doc.save("C:\\Users\\PRO\\OneDrive\\Documents\\GitHub\\pega\\Pega_Movie_Ticket_Booking_Report.docx")
    doc.save("C:\\Users\\PRO\\OneDrive\\Documents\\GitHub\\pega\\RD_Sourav.docx")
    print("Report generated successfully as both Pega_Movie_Ticket_Booking_Report.docx and RD_Sourav.docx!")

if __name__ == "__main__":
    generate_report()
